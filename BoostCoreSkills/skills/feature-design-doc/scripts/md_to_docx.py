#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert a design-doc Markdown into a .docx that honors a fixed corporate
format, embedding the pre-rendered figure PNGs that the Markdown references.

Format constraints (do not loosen without being asked):
  * 中文正文     -> 宋体, 五号 (10.5pt)
  * 英文/数字正文 -> Times New Roman, 五号 (10.5pt)
  * 嵌入正文的 `代码词` -> 与正文同字体 (统一观感),不单独用等宽
  * 完整代码块   -> Consolas 9pt, 浅灰底单元格
  * 字体一律自动黑色 (无彩色),不用斜体
  * 段落单倍行距,段后 6 磅

Markdown features handled: ATX headings, GitHub pipe tables, fenced code blocks
(``` ... ```), blockquotes (>), bullet/ordered lists, inline **bold** / `code`,
horizontal rules (---), and image refs ![cap](images/figNN.png). Image paths are
resolved relative to the Markdown file's directory.

Usage:
  python md_to_docx.py <doc.md> [out.docx] [--images-base DIR]

Requires: python-docx, pillow  (see setup_env.sh).
"""
import argparse, os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

CJK = "宋体"
LATIN = "Times New Roman"
MONO = "Consolas"
BODY_PT = 10.5
BLACK = RGBColor(0, 0, 0)
MAX_W = 6.3  # inches

def build(md_path, out_path, images_base):
    md_dir = os.path.dirname(os.path.abspath(md_path))
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = LATIN
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
    pf = normal.paragraph_format
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0); pf.space_after = Pt(6)

    def set_ea(run, font, latin=None):
        latin = latin or font
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rpr.append(rf)
        rf.set(qn("w:eastAsia"), font); rf.set(qn("w:ascii"), latin); rf.set(qn("w:hAnsi"), latin)

    def style(run, bold=False, mono=False, size=None):
        if mono:
            run.font.name = MONO; set_ea(run, CJK, latin=MONO)
        else:
            run.font.name = LATIN; set_ea(run, CJK, latin=LATIN)
        run.italic = False
        if bold: run.bold = True
        run.font.size = Pt(size if size else BODY_PT)
        run.font.color.rgb = BLACK

    def inline(p, text):
        for tok in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
            if not tok: continue
            if tok.startswith("**") and tok.endswith("**"):
                style(p.add_run(tok[2:-2]), bold=True)
            elif tok.startswith("`") and tok.endswith("`"):
                style(p.add_run(tok[1:-1]))          # 嵌入代码词 = 正文字体
            else:
                style(p.add_run(tok))

    def shade(cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hexc)
        tcPr.append(sh)

    def code_block(lines):
        t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.cell(0, 0); shade(cell, "F4F6F8"); cell.width = Inches(MAX_W)
        tcPr = cell._tc.get_or_add_tcPr(); b = OxmlElement("w:tcBorders")
        for e in ("top","left","bottom","right"):
            el = OxmlElement(f"w:{e}"); el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4"); el.set(qn("w:color"),"D0D5DB"); b.append(el)
        tcPr.append(b)
        para = cell.paragraphs[0]; para.paragraph_format.space_after = Pt(0); para.paragraph_format.space_before = Pt(0)
        for i, ln in enumerate(lines):
            style(para.add_run(("" if i == 0 else "\n") + ln), mono=True, size=9)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def add_image(rel, cap):
        png = rel if os.path.isabs(rel) else os.path.join(md_dir, rel)
        if not os.path.exists(png) and images_base:
            png = os.path.join(images_base, os.path.basename(rel))
        if not os.path.exists(png):
            print("WARN missing image:", rel); return
        w, _ = Image.open(png).size
        disp = min(MAX_W, w / 192.0)
        if disp < 2.6: disp = min(MAX_W, w / 130.0)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(png, width=Inches(disp))
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style(c.add_run(cap), size=9); c.paragraph_format.space_after = Pt(6)

    def flush_table(buf):
        rows = [[c.strip() for c in ln.strip().strip("|").split("|")] for ln in buf]
        rows = [r for r in rows if not all(set(c) <= set("-: ") and c for c in r)]
        if not rows: return
        nc = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=nc); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(rows):
            for ci in range(nc):
                cell = t.cell(ri, ci); cell.paragraphs[0].text = ""
                if ri == 0: shade(cell, "EDEDED")
                inline(cell.paragraphs[0], row[ci] if ci < len(row) else "")
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9.5)
                    if ri == 0: run.bold = True
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    src = open(md_path, encoding="utf-8").read().splitlines()
    i = 0; in_code = False; code = []; tbl = []
    def flush_tbl():
        nonlocal tbl
        if tbl: flush_table(tbl); tbl = []

    while i < len(src):
        line = src[i]
        m = re.match(r"^```(\w*)\s*$", line)
        if m and not in_code:
            flush_tbl(); in_code = True; code = []; i += 1; continue
        if line.strip() == "```" and in_code:
            in_code = False; code_block(code); i += 1; continue
        if in_code:
            code.append(line); i += 1; continue

        im = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", line)
        if im:
            flush_tbl(); add_image(im.group(2), im.group(1)); i += 1; continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            tbl.append(line); i += 1; continue
        else:
            flush_tbl()

        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            lvl = len(h.group(1))
            hd = doc.add_heading(level=min(lvl, 4)); hd.text = ""
            r = hd.add_run(h.group(2).strip()); set_ea(r, CJK, latin=LATIN)
            r.font.color.rgb = BLACK; r.bold = True; r.italic = False
            r.font.size = Pt({1:16,2:14,3:12,4:11}.get(min(lvl,4),11))
            hd.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            hd.paragraph_format.space_before = Pt(6); hd.paragraph_format.space_after = Pt(6)
            i += 1; continue

        if line.strip() == "---":
            p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
            bd = OxmlElement("w:pBdr"); bo = OxmlElement("w:bottom")
            bo.set(qn("w:val"),"single"); bo.set(qn("w:sz"),"6"); bo.set(qn("w:color"),"C0C0C0"); bo.set(qn("w:space"),"1")
            bd.append(bo); pPr.append(bd); i += 1; continue

        if line.startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25)
            style(p.add_run("▍ ")); inline(p, line.lstrip(">").strip()); i += 1; continue

        bm = re.match(r"^(\s*)([-*])\s+(.*)$", line); om = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if bm:
            inline(doc.add_paragraph(style="List Bullet"), bm.group(3)); i += 1; continue
        if om:
            inline(doc.add_paragraph(style="List Number"), om.group(3)); i += 1; continue

        if line.strip() == "":
            i += 1; continue
        inline(doc.add_paragraph(), line); i += 1

    flush_tbl()
    doc.save(out_path)
    print("saved", out_path)

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("md")
    ap.add_argument("out", nargs="?", default=None)
    ap.add_argument("--images-base", default=None)
    a = ap.parse_args()
    out = a.out or os.path.splitext(a.md)[0] + ".docx"
    build(a.md, out, a.images_base)
